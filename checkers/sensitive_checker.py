import os
import re
import platform
from concurrent.futures import ThreadPoolExecutor, as_completed
from checkers.base_checker import BaseChecker, CheckResult


class SensitiveChecker(BaseChecker):

    ID_PATTERN = re.compile(r"\b\d{17}[\dXx]\b")
    PHONE_PATTERN = re.compile(r"\b1[3-9]\d{9}\b")
    EMAIL_PATTERN = re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")

    def __init__(self, extra_dirs: list = None):
        super().__init__()
        self.extra_dirs = extra_dirs or []

    def run(self) -> CheckResult:
        self.logger.info("SensitiveChecker started")
        dirs = self._get_scan_dirs()
        keywords = self.config.get("sensitive_keywords", [])
        extensions = set(self.config.get("scan_extensions", []))
        max_mb = self.config.get("max_file_size_mb", 50)
        max_bytes = max_mb * 1024 * 1024
        self.logger.info(f"  Scan dirs  : {dirs}")
        self.logger.info(f"  Extensions : {sorted(extensions)}, max size: {max_mb}MB")

        files = []
        for d in dirs:
            files.extend(self._collect_files(d, extensions, max_bytes))
        self.logger.info(f"  Files to scan: {len(files)}")

        violations = []
        with ThreadPoolExecutor(max_workers=4) as pool:
            futures = {
                pool.submit(self._scan_file, f, keywords, max_bytes): f for f in files
            }
            for future in as_completed(futures):
                result = future.result()
                if result:
                    violations.append(result)

        passed = len(violations) == 0
        total_scanned = len(files)

        summary = f"扫描目录: {dirs}\n扫描文件总数: {total_scanned}\n发现敏感文件数: {len(violations)}\n"
        if violations:
            summary += "\n敏感文件列表:\n"
            for v in violations:
                summary += f"  - {v['path']}\n    命中: {', '.join(v['hits'])}\n"
        else:
            summary += "\n未发现敏感信息文件。"

        recommendations = []
        if violations:
            recommendations.append("请及时清理或加密上述包含敏感信息的文件，避免明文存储身份证号、手机号、密级关键词等信息。")

        self.logger.info(f"SensitiveChecker done: {len(files)} scanned, {len(violations)} sensitive files found")
        for v in violations:
            self.logger.info(f"  [SENSITIVE] {v['path']} — hits: {v['hits']}")
        return CheckResult(
            module="敏感信息检查",
            passed=passed,
            violations=violations,
            summary=summary,
            recommendations=recommendations,
        )

    def _get_scan_dirs(self) -> list:
        home = os.path.expanduser("~")
        system = platform.system()
        if system == "Windows":
            defaults = [
                os.path.join(home, "Desktop"),
                os.path.join(home, "Documents"),
            ]
        else:
            defaults = [
                os.path.join(home, "Desktop"),
                os.path.join(home, "Documents"),
            ]
        configured = self.config.get("default_scan_dirs", [])
        all_dirs = defaults + configured + self.extra_dirs
        return [d for d in all_dirs if os.path.isdir(d)]

    def _collect_files(self, directory: str, extensions: set, max_bytes: int) -> list:
        result = []
        try:
            for entry in os.scandir(directory):
                if entry.is_dir(follow_symlinks=False):
                    result.extend(self._collect_files(entry.path, extensions, max_bytes))
                elif entry.is_file(follow_symlinks=False):
                    _, ext = os.path.splitext(entry.name)
                    if ext.lower() in extensions:
                        try:
                            if entry.stat().st_size <= max_bytes:
                                result.append(entry.path)
                        except OSError:
                            pass
        except PermissionError:
            pass
        except Exception as e:
            self.logger.warning(f"Error scanning directory: {directory}", exc_info=True)
        return result

    def _scan_file(self, path: str, keywords: list, max_bytes: int):
        hits = []
        _, ext = os.path.splitext(path)
        ext = ext.lower()
        try:
            if ext in (".txt", ".csv", ".md"):
                text = self._read_text_partial(path)
            elif ext == ".docx":
                text = self._read_docx(path)
            elif ext == ".doc":
                text = self._read_doc(path)
            elif ext == ".pdf":
                text = self._read_pdf(path)
            elif ext in (".xls", ".xlsx"):
                text = self._read_excel(path)
            else:
                return None

            if self.ID_PATTERN.search(text):
                hits.append("身份证号")
            if self.PHONE_PATTERN.search(text):
                hits.append("手机号")
            if self.EMAIL_PATTERN.search(text):
                hits.append("邮箱地址")
            for kw in keywords:
                if kw in text:
                    hits.append(f"密级关键词({kw})")

        except Exception as e:
            self.logger.debug(f"Could not scan file: {path}", exc_info=True)
            return None

        if hits:
            return {"path": path, "hits": hits}
        return None

    def _read_text_partial(self, path: str) -> str:
        chunk = 10 * 1024
        try:
            with open(path, "rb") as f:
                head = f.read(chunk)
                try:
                    f.seek(-chunk, 2)
                except OSError:
                    f.seek(0)
                tail = f.read(chunk)
            return (head + tail).decode("utf-8", errors="ignore")
        except Exception:
            return ""

    def _read_docx(self, path: str) -> str:
        """Read .docx (Office Open XML) format."""
        try:
            import docx
            doc = docx.Document(path)
            texts = [p.text for p in doc.paragraphs[:200]]
            # Also scan table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        texts.append(cell.text)
            return "\n".join(texts)
        except Exception:
            return ""

    def _read_doc(self, path: str) -> str:
        """Read old .doc (binary Word 97-2003) format."""
        # Try win32com on Windows (requires Word installed)
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(os.path.abspath(path))
            text = doc.Content.Text
            doc.Close(False)
            word.Quit()
            return text
        except Exception:
            pass

        # Try antiword (Linux / macOS)
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", path], capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0 and result.stdout:
                return result.stdout
        except Exception:
            pass

        # Fallback: extract printable ASCII + attempt UTF-16-LE decode
        try:
            with open(path, "rb") as f:
                data = f.read(500 * 1024)
            ascii_parts = re.findall(rb"[\x20-\x7e]{4,}", data)
            ascii_text = b" ".join(ascii_parts).decode("ascii", errors="ignore")
            try:
                utf16_text = data.decode("utf-16-le", errors="ignore")
            except Exception:
                utf16_text = ""
            return ascii_text + "\n" + utf16_text
        except Exception:
            return ""

    def _read_pdf(self, path: str) -> str:
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages[:10])
        except Exception:
            return ""

    def _read_excel(self, path: str) -> str:
        """Read all sheets from .xls or .xlsx, up to 500 rows each."""
        try:
            import pandas as pd
            _, ext = os.path.splitext(path)
            engine = "xlrd" if ext.lower() == ".xls" else None
            xl = pd.ExcelFile(path, engine=engine)
            texts = []
            for sheet in xl.sheet_names:
                try:
                    df = xl.parse(sheet, nrows=500)
                    texts.append(df.to_string())
                except Exception:
                    continue
            return "\n".join(texts)
        except Exception:
            return ""
